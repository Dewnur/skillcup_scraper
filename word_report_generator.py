import re

import docx
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Mm
from docx.shared import Pt

from db_models_manager import *
from models import *
from tools.text_preprocessor import comment_tokens


def _split_text_and_link(s: str):
    pattern = r'(https?://[^\s]+)'
    result = re.split(pattern, s)
    return result


class WordReportGenerator:

    def __init__(self, path_dir: str, cards: List[Card]):
        self.path_dir = path_dir
        self._doc = Document()
        self.cards = sorted(cards, key=lambda card: card.sequence_number)
        self._filename_extension = '.docx'

    def _init_styles(self):
        header_style = self._doc.styles.add_style('Header1', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = 'Times New Roman'
        header_style.font.size = Pt(11)
        header_style.font.bold = True
        header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        header_style.paragraph_format.left_indent = Mm(-3)
        title_style = self._doc.styles.add_style('title_style', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(12)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        task_name_style = self._doc.styles.add_style('task_name_style', WD_STYLE_TYPE.PARAGRAPH)
        task_name_style.font.name = 'Roboto'
        task_name_style.font.size = Pt(11)
        task_name_style.font.bold = True
        task_name_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        comment_style = self._doc.styles.add_style('comment_style', WD_STYLE_TYPE.PARAGRAPH)
        comment_style.font.name = 'Arial'
        comment_style.font.size = Pt(11)
        comment_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        comment_style.paragraph_format.left_indent = Mm(-10)

    def generate_report(self, person: Person):
        self._init_styles()
        header = self._doc.add_paragraph(f'ФИ: {person.name}\nТариф: {person.rate}\nСсылка на канал: ', style='Header1')
        self._add_hyperlink(header, text=f'{person.tg_url}', url=f'{person.tg_url}', color='0000FF', underline=False)
        punc = '________'
        header.add_run(f'\n{punc * 10}')
        save_path = ''
        for card in self.cards:
            self._doc.add_paragraph(f'ДОМАШКА {card.sequence_number}:', style='title_style')
            tasks = fetchall(Task, card_id=card.id)
            sorted_tasks = sorted(tasks, key=lambda task: task.sequence_number)
            for task in sorted_tasks:
                self._doc.add_paragraph(f'{task.name}:', style='task_name_style')
                comments = fetchall(Comment, task_id=task.id, person_id=person.id)
                if comments is None:
                    self._doc.add_paragraph(f'Нет ответа', style='comment_style')
                    continue
                sorted_comments = sorted(comments, key=lambda comment: comment.sequence_number)
                content_list = [comment.content for comment in sorted_comments]
                content = '\n'.join(content_list)
                content_tokens = comment_tokens(content)
                comment_paragraph = self._doc.add_paragraph(f'', style='comment_style')
                for substring, link in content_tokens:
                    if link:
                        self._add_hyperlink(
                            comment_paragraph,
                            text=f'{substring}',
                            url=f'{substring}',
                            color='0000FF',
                            underline=False
                        )
                    else:
                        comment_paragraph.add_run(f'{substring}')
            person_card = fetchone(PersonCard, person_id=person.id, card_id=card.id)
            save_path = self.path_dir + f'{person.name}'
            if person_card.total_done is not None:
                save_path += f' {person_card.total_done}'
            save_path += self._filename_extension

        self._doc.save(save_path)

    def _add_hyperlink(self, paragraph, text, url, color, underline):
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if not color is None:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        # Remove underlining if it is requested
        if not underline:
            u = docx.oxml.shared.OxmlElement('w:u')
            u.set(docx.oxml.shared.qn('w:val'), 'single')
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        return hyperlink
